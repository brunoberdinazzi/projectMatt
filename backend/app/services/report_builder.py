from __future__ import annotations

import os
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from ..models import ReportBuildRequest
from .report_terms import entity_display_name


class ReportBuilder:
    def build(
        self,
        payload: ReportBuildRequest,
        output_format: str = "docx",
        template_path: Optional[Path] = None,
    ) -> Path:
        if output_format == "pdf":
            return self._build_pdf(payload)
        return self._build_docx(payload, template_path=template_path)

    def _build_docx(self, payload: ReportBuildRequest, template_path: Optional[Path] = None) -> Path:
        resolved_template = template_path or self._default_template_path()
        document = Document(str(resolved_template)) if resolved_template else Document()
        if resolved_template:
            self._clear_document_body(document)
            self._populate_template_report(document, payload)
        else:
            self._populate_plain_report(document, payload)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            output_path = Path(temp_file.name)

        document.save(output_path)
        return output_path

    def _populate_plain_report(self, document: Document, payload: ReportBuildRequest) -> None:
        document.add_heading(payload.titulo_relatorio, level=0)

        if payload.orgao:
            document.add_paragraph(f"Entidade: {payload.orgao}")
        if payload.tipo_orgao:
            document.add_paragraph(f"Tipo de entidade: {payload.tipo_orgao}")
        if payload.periodo_analise:
            document.add_paragraph(f"Periodo da analise: {payload.periodo_analise}")

        for section in payload.secoes:
            document.add_heading(section.titulo, level=1)
            self._add_section_content(document, section)

    def _populate_template_report(self, document: Document, payload: ReportBuildRequest) -> None:
        self._add_cover_page(document, payload)
        document.add_page_break()

        for section in payload.secoes:
            self._add_heading(document, section.titulo)
            self._add_section_content(document, section)

    def _add_cover_page(self, document: Document, payload: ReportBuildRequest) -> None:
        title = document.add_paragraph(style="Normal")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("RELATORIO TECNICO")
        run.bold = True
        run.font.size = Pt(15)

        if payload.numero_relatorio:
            number = document.add_paragraph(style="Normal")
            number.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = number.add_run(payload.numero_relatorio)
            run.bold = True
            run.font.size = Pt(11)

        for line in (
            f"Solicitante: {payload.promotoria}" if payload.promotoria else None,
            f"Referencia: {payload.referencia}" if payload.referencia else None,
            f"Solicitacao: {payload.solicitacao}" if payload.solicitacao else None,
        ):
            if not line:
                continue
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run(line)

        document.add_paragraph("")

        subtitle = document.add_paragraph(style="Normal")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(payload.titulo_relatorio.upper())
        run.bold = True
        run.font.size = Pt(12)

        cover_orgao = self._cover_orgao_label(payload)
        if cover_orgao:
            orgao = document.add_paragraph(style="Normal")
            orgao.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = orgao.add_run(cover_orgao)
            run.bold = True
            run.font.size = Pt(12)

        for line in (payload.cidade_emissao, self._cover_period_label(payload.data_emissao)):
            if not line:
                continue
            paragraph = document.add_paragraph(style="Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            run.font.size = Pt(11)

        document.add_paragraph("")

    def _add_heading(self, document: Document, text: str) -> None:
        style_name = "Heading 1" if text == text.upper() else "Heading 2"
        paragraph = document.add_paragraph(style=style_name)
        paragraph.add_run(text)

    def _add_body_paragraph(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph(style="Body Text")
        paragraph.add_run(text)

    def _add_section_content(self, document: Document, section) -> None:
        if section.texto:
            for paragraph in section.texto.split("\n\n"):
                self._add_structured_block(document, paragraph)
        if section.table_headers and section.table_rows:
            self._add_table(document, section.table_headers, section.table_rows)
            document.add_paragraph("")

    def _add_structured_block(self, document: Document, text: str) -> None:
        cleaned = text.strip()
        if not cleaned:
            return
        lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
        if lines and all(line.startswith("- ") for line in lines):
            for line in lines:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(line[2:])
            return
        self._add_body_paragraph(document, cleaned)

    def _add_table(self, document: Document, headers: list[str], rows: list[list[str]]) -> None:
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            header_cells[index].text = header
            if header_cells[index].paragraphs and header_cells[index].paragraphs[0].runs:
                header_cells[index].paragraphs[0].runs[0].bold = True

        for row in rows:
            row_cells = table.add_row().cells
            for index, value in enumerate(row[: len(headers)]):
                row_cells[index].text = value or "-"

    def _clear_document_body(self, document: Document) -> None:
        body = document._element.body
        for child in list(body):
            if child.tag.endswith("sectPr"):
                continue
            body.remove(child)

    def _default_template_path(self) -> Optional[Path]:
        value = os.getenv("REPORT_TEMPLATE_PATH")
        if not value:
            return None
        path = Path(value).expanduser()
        return path if path.exists() else None

    def _cover_orgao_label(self, payload: ReportBuildRequest) -> str:
        return entity_display_name(payload.orgao, payload.tipo_orgao).replace("entidade nao informada", "")

    def _cover_period_label(self, date_text: Optional[str]) -> Optional[str]:
        if not date_text:
            return None
        normalized = " ".join(date_text.lower().split())
        month_map = {
            "janeiro": "Janeiro",
            "fevereiro": "Fevereiro",
            "marco": "Marco",
            "abril": "Abril",
            "maio": "Maio",
            "junho": "Junho",
            "julho": "Julho",
            "agosto": "Agosto",
            "setembro": "Setembro",
            "outubro": "Outubro",
            "novembro": "Novembro",
            "dezembro": "Dezembro",
        }
        for month, label in month_map.items():
            token = f" de {month} de "
            if token in normalized:
                year = normalized.split(token, 1)[1].strip()
                if year:
                    return f"{label}/{year}"
        return date_text

    def _build_pdf(self, payload: ReportBuildRequest) -> Path:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            output_path = Path(temp_file.name)

        styles = self._pdf_styles()
        story = [Paragraph(payload.titulo_relatorio, styles["DocTitle"]), Spacer(1, 0.35 * cm)]

        metadata_lines = [
            f"Entidade: {payload.orgao}" if payload.orgao else None,
            f"Tipo de entidade: {payload.tipo_orgao}" if payload.tipo_orgao else None,
            f"Periodo da analise: {payload.periodo_analise}" if payload.periodo_analise else None,
        ]
        for line in metadata_lines:
            if line:
                story.append(Paragraph(self._escape(line), styles["DocBody"]))
                story.append(Spacer(1, 0.12 * cm))

        if any(metadata_lines):
            story.append(Spacer(1, 0.18 * cm))

        for section in payload.secoes:
            story.append(Paragraph(self._escape(section.titulo), styles["DocHeading"]))
            story.append(Spacer(1, 0.12 * cm))
            for paragraph in section.texto.split("\n\n"):
                for line in paragraph.splitlines():
                    story.append(Paragraph(self._escape(line), styles["DocBody"]))
                    story.append(Spacer(1, 0.08 * cm))
                story.append(Spacer(1, 0.12 * cm))
            if section.table_headers and section.table_rows:
                story.append(Paragraph(self._escape(" | ".join(section.table_headers)), styles["DocBody"]))
                story.append(Spacer(1, 0.08 * cm))
                for row in section.table_rows:
                    story.append(Paragraph(self._escape(" | ".join(row)), styles["DocBody"]))
                    story.append(Spacer(1, 0.06 * cm))
                story.append(Spacer(1, 0.14 * cm))

        document = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=2.2 * cm,
            rightMargin=2.2 * cm,
            topMargin=2.0 * cm,
            bottomMargin=2.0 * cm,
            title=payload.titulo_relatorio,
        )
        document.build(story)
        return output_path

    def _pdf_styles(self):
        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                name="DocTitle",
                parent=styles["Title"],
                fontName="Helvetica-Bold",
                fontSize=18,
                leading=22,
                textColor=HexColor("#0F172A"),
            )
        )
        styles.add(
            ParagraphStyle(
                name="DocHeading",
                parent=styles["Heading2"],
                fontName="Helvetica-Bold",
                fontSize=12,
                leading=15,
                textColor=HexColor("#1D4ED8"),
                spaceBefore=6,
            )
        )
        styles.add(
            ParagraphStyle(
                name="DocBody",
                parent=styles["BodyText"],
                fontName="Helvetica",
                fontSize=10.2,
                leading=14,
                textColor=HexColor("#111827"),
            )
        )
        return styles

    def _escape(self, text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
